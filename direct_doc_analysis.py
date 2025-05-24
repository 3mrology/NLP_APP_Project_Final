import os
import sys
import re

# Try importing required packages and install if not available
try:
    from docx import Document
except ImportError:
    print("python-docx not available. Installing...")
    import subprocess
    subprocess.call([sys.executable, "-m", "pip", "install", "python-docx"])
    try:
        from docx import Document
    except ImportError:
        print("Failed to import docx even after installation.")

try:
    import PyPDF2
except ImportError:
    print("PyPDF2 not available. Installing...")
    import subprocess
    subprocess.call([sys.executable, "-m", "pip", "install", "PyPDF2"])
    try:
        import PyPDF2
    except ImportError:
        print("Failed to import PyPDF2 even after installation.")

def analyze_docx(file_path):
    """Analyze a DOCX file and print its content"""
    print(f"\nAnalyzing DOCX: {os.path.basename(file_path)}")
    print("-" * 80)
    
    try:
        # Load the document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Display document info
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total sections: {len(doc.sections)}")
        
        # Print document headings
        print("\nDocument Structure:")
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading'):
                print(f"{para.style.name}: {para.text}")
        
        # Print first few paragraphs for context
        print("\nDocument Preview:")
        preview_text = "\n".join(full_text[:10])
        print(preview_text)
        
        # Try to extract key entities
        entities = extract_entities(full_text)
        print("\nKey Entities:")
        for category, items in entities.items():
            if items:
                print(f"{category}: {', '.join(items[:10])}")
        
        return full_text
    
    except Exception as e:
        print(f"Error analyzing DOCX: {str(e)}")
        return []

def analyze_pdf_metadata(file_path):
    """Analyze basic metadata of a PDF file"""
    print(f"\nAnalyzing PDF: {os.path.basename(file_path)}")
    print("-" * 80)
    
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            # Get number of pages
            num_pages = len(reader.pages)
            print(f"Number of pages: {num_pages}")
            
            # Get document metadata
            if reader.metadata:
                print("\nDocument Metadata:")
                for key, value in reader.metadata.items():
                    if isinstance(key, str) and key.startswith('/'):
                        clean_key = key[1:]  # Remove leading slash
                        print(f"{clean_key}: {value}")
            
            # Extract text from first few pages
            print("\nDocument Preview (first page):")
            if num_pages > 0:
                text = reader.pages[0].extract_text()
                print(text[:1000] + "..." if len(text) > 1000 else text)
                
                # Try to extract potential entities from first few pages
                all_text = ""
                for i in range(min(3, num_pages)):
                    all_text += reader.pages[i].extract_text() + " "
                
                entities = extract_entities([all_text])
                print("\nKey Entities:")
                for category, items in entities.items():
                    if items:
                        print(f"{category}: {', '.join(items[:10])}")
                        
    except ImportError:
        print("PyPDF2 not available. Installing...")
        import subprocess
        subprocess.call([sys.executable, "-m", "pip", "install", "PyPDF2"])
        print("Please run this script again to analyze the PDF.")
    except Exception as e:
        print(f"Error analyzing PDF: {str(e)}")

def extract_entities(text_blocks):
    """Extract potential entities from text"""
    text = " ".join(text_blocks)
    
    # Initialize entity categories
    entities = {
        "People": set(),
        "Disorders": set(),
        "Therapies": set(),
        "Technologies": set()
    }
    
    # Define patterns for each category
    person_pattern = r'[A-Z][a-z]+ [A-Z][a-z]+'
    disorder_patterns = [
        r'(?:Major )?Depressive Disorder',
        r'Generalized Anxiety Disorder',
        r'Bipolar Disorder',
        r'Panic Disorder',
        r'PTSD',
        r'Post-Traumatic Stress Disorder',
        r'Schizophrenia',
        r'OCD',
        r'Obsessive-Compulsive Disorder'
    ]
    therapy_patterns = [
        r'Cognitive[ -]Behavioral Therapy',
        r'CBT',
        r'Reminiscence Therapy',
        r'Psychodynamic Therapy',
        r'Exposure Therapy',
        r'Dialectical[ -]Behavior Therapy',
        r'DBT',
        r'Role[ -]Playing Therapy'
    ]
    technology_patterns = [
        r'Large Language Model',
        r'LLM',
        r'Knowledge Graph',
        r'Neural Network',
        r'Deep Learning',
        r'Machine Learning',
        r'NLP',
        r'Natural Language Processing',
        r'AI',
        r'Artificial Intelligence',
        r'GPT',
        r'BERT',
        r'Llama',
        r'Vector Database',
        r'Embedding',
        r'RAG',
        r'Retrieval[ -]Augmented Generation',
        r'Azure OpenAI',
        r'OpenAI',
        r'HuggingFace',
        r'Neo4j'
    ]
    
    # Find people
    for match in re.finditer(person_pattern, text):
        name = match.group(0)
        if len(name.split()) == 2:  # Ensure it's first and last name
            entities["People"].add(name)
    
    # Find disorders
    for pattern in disorder_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            entities["Disorders"].add(match.group(0))
    
    # Find therapies
    for pattern in therapy_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            entities["Therapies"].add(match.group(0))
    
    # Find technologies
    for pattern in technology_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            entities["Technologies"].add(match.group(0))
    
    return {k: list(v) for k, v in entities.items()}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python direct_doc_analysis.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
    
    # Analyze based on file extension
    if file_path.lower().endswith('.docx'):
        analyze_docx(file_path)
    elif file_path.lower().endswith('.pdf'):
        analyze_pdf_metadata(file_path)
    else:
        print(f"Unsupported file format: {os.path.splitext(file_path)[1]}")
